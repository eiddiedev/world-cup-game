from __future__ import annotations

import hashlib
import itertools
import math
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from parse_2026_rosters import URL as ROSTER_URL, parse


OUTPUT = Path("/Users/a1234/Desktop/剑指美加墨_16队24人球员数据与名单替换稿_2026-07-22.docx")

AWARDS_URL = "https://www.fifa.com/en/tournaments/mens/worldcup/canadamexicousa2026/articles/award-winners"
ASSISTS_URL = "https://www.fifa.com/en/tournaments/mens/worldcup/canadamexicousa2026/articles/most-assists-top-assisters"
RESULTS_URL = "https://www.fifa.com/en/articles/knockout-stage-match-schedule-bracket"
POWER_URL = "https://www.fifa.com/en/tournaments/mens/worldcup/canadamexicousa2026/power-rankings"

TEAM_ORDER = [
    "西班牙", "阿根廷", "法国", "英格兰", "巴西", "葡萄牙", "德国", "日本",
    "摩洛哥", "挪威", "哥伦比亚", "美国", "加拿大", "墨西哥", "佛得角", "库拉索",
]

POSITION_MAP = {"门将": "GK", "后卫": "DF", "中场": "MF", "前锋": "FW"}

DROP = {
    "西班牙": ["霍安·加西亚", "博尔哈·伊格莱西亚斯"],
    "阿根廷": ["胡安·穆索"],
    "法国": ["罗班·里塞", "马克桑斯·拉克鲁瓦"],
    "英格兰": ["詹姆斯·特拉福德", "贾雷尔·宽萨"],
    "巴西": ["维弗顿", "拉扬"],
    "葡萄牙": ["鲁伊·席尔瓦", "贡萨洛·格德斯"],
    "德国": ["亚历山大·尼贝尔", "纳迪姆·阿米里"],
    "日本": ["早川友基", "盐贝健人"],
    "摩洛哥": ["塔纳乌蒂", "萨拉赫-埃丁"],
    "挪威": ["桑德·唐维克", "亨里克·法尔克纳"],
    "哥伦比亚": ["阿尔瓦罗·蒙特罗", "安德烈斯·戈麦斯"],
    "美国": ["克里斯·布雷迪", "曾德哈斯"],
    "加拿大": ["欧文·古德曼", "杰登·纳尔逊"],
    "墨西哥": ["卡洛斯·阿塞韦多", "吉列尔莫·马丁内斯"],
    "佛得角": ["CJ·多斯桑托斯", "斯托皮拉"],
    "库拉索": ["特雷弗·多恩布施", "德弗龙·方维尔"],
}


@dataclass(frozen=True)
class TeamMeta:
    prefix: str
    result: str
    difficulty: str
    jersey: str
    skill: str
    personality: str
    delta: int


TEAM_META = {
    "西班牙": TeamMeta("红潮", "冠军", "★", "红色主场 / 浅色客场", "控场王冠", "以控球、压迫和阵地耐心统治比赛", 3),
    "阿根廷": TeamMeta("蓝白", "亚军", "★", "蓝白条纹 / 深色客场", "绝境反击", "越接近终场，核心球员越危险", 3),
    "法国": TeamMeta("高卢", "四强", "★", "深蓝主场 / 白色客场", "纵深爆破", "锋线速度和单点爆发决定上限", 3),
    "英格兰": TeamMeta("三狮", "季军", "★", "白色主场 / 深蓝客场", "二次进攻", "身体、定位球与禁区持续施压", 3),
    "巴西": TeamMeta("桑巴", "十六强", "★★", "黄色主场 / 蓝色客场", "桑巴节奏", "单点技术出众，强调边路一对一", 1),
    "葡萄牙": TeamMeta("葡国", "十六强", "★★", "红色主场 / 白色客场", "关键先生", "老将终结与中场创造力并存", 1),
    "德国": TeamMeta("德意", "三十二强", "★★★", "白色主场 / 深色客场", "日耳曼机器", "结构完整，重视执行和高强度跑动", 0),
    "日本": TeamMeta("蓝武", "三十二强", "★★★", "深蓝主场 / 浅色客场", "高压逼抢", "快速传切与团队协作弥补对抗差距", -1),
    "摩洛哥": TeamMeta("北非", "八强", "★★", "红色主场 / 白色客场", "沙漠反击", "防线韧性与边路推进形成强反击", 2),
    "挪威": TeamMeta("北海", "八强", "★★", "红色主场 / 白色客场", "北欧巨人", "中锋终结与核心输送构成最短进球路径", 2),
    "哥伦比亚": TeamMeta("咖啡", "十六强", "★★★", "黄色主场 / 深色客场", "咖啡旋律", "边路冲击和前场创造力制造机会", 0),
    "美国": TeamMeta("星条", "十六强", "★★★", "白色主场 / 深蓝客场", "主场浪潮", "主场节奏、跑动和纵向推进形成压迫", 0),
    "加拿大": TeamMeta("枫叶", "十六强", "★★★", "红色主场 / 白色客场", "枫叶快攻", "速度型边路和快速转换是主要武器", 0),
    "墨西哥": TeamMeta("绿鹰", "十六强", "★★★", "绿色主场 / 白色客场", "高原节奏", "中前场连续跑动和主场气氛强化攻势", 0),
    "佛得角": TeamMeta("蓝鲨", "三十二强", "★★★★", "蓝色主场 / 白色客场", "蓝鲨冲刺", "对抗、纵向推进和团队纪律制造爆冷", -4),
    "库拉索": TeamMeta("海岛", "小组赛", "★★★★★", "蓝黄主场 / 白色客场", "海岛之心", "门将发挥和反击效率决定弱队上限", -7),
}


@dataclass(frozen=True)
class StarSpec:
    keyword: str
    alias: str
    tier: str
    stats: tuple[int, int, int, int, int, int]
    note: str


STARS = {
    "西班牙": [
        StarSpec("罗德里", "红潮司令", "gold", (76, 88, 96, 94, 92, 5), "金球奖；攻防节拍与控场核心"),
        StarSpec("乌奈·西蒙", "红潮门神", "silver", (60, 86, 88, 97, 92, 5), "金手套；8场7次零封"),
        StarSpec("库巴西", "红潮新墙", "silver", (82, 82, 88, 93, 88, 5), "最佳年轻球员；高位防线出球核心"),
        StarSpec("亚马尔", "红潮神童", "silver", (94, 72, 94, 48, 90, 5), "冠军阵容边路爆点与创造者"),
    ],
    "阿根廷": [
        StarSpec("梅西", "潘帕球王", "gold", (82, 70, 99, 45, 80, 5), "银球奖、银靴；淘汰赛关键创造者"),
        StarSpec("埃米利亚诺", "潘帕门神", "silver", (56, 88, 80, 93, 89, 5), "大赛型门将，点球与关键扑救加成"),
        StarSpec("阿尔瓦雷斯", "蛛网猎手", "silver", (90, 82, 91, 52, 94, 5), "前场压迫、跑位和终结兼备"),
        StarSpec("麦卡利斯特", "蓝白节拍", "silver", (80, 82, 90, 82, 92, 4), "中场推进与攻防连接核心"),
    ],
    "法国": [
        StarSpec("姆巴佩", "高卢闪电", "gold", (99, 84, 96, 42, 94, 5), "铜球奖、金靴；8场10球"),
        StarSpec("奥利塞", "高卢画师", "silver", (91, 75, 95, 55, 92, 5), "赛事5次助攻，创造力榜首"),
        StarSpec("登贝莱", "双翼魔术", "silver", (95, 72, 92, 48, 87, 5), "双足边路爆破与2次助攻"),
    ],
    "英格兰": [
        StarSpec("哈里·凯恩", "三狮重炮", "gold", (82, 90, 92, 45, 88, 5), "支点、组织和禁区终结兼备"),
        StarSpec("贝林厄姆", "三狮帝星", "silver", (88, 92, 94, 84, 95, 5), "铜靴；中场推进与禁区终结核心"),
        StarSpec("皮克福德", "三狮门神", "silver", (58, 84, 87, 94, 91, 5), "赛事控球阶段表现突出的门将"),
        StarSpec("萨卡", "三狮飞翼", "silver", (92, 78, 91, 55, 92, 5), "边路核心，赛事2次助攻"),
    ],
    "巴西": [
        StarSpec("内马尔", "桑巴魔术", "gold", (80, 70, 92, 40, 72, 5), "技术、关键传球和关键球上限极高"),
        StarSpec("维尼修斯", "桑巴飞刃", "silver", (96, 77, 92, 42, 89, 5), "边路最强爆点与转换终结者"),
        StarSpec("吉马良斯", "桑巴节拍", "silver", (82, 86, 92, 84, 94, 5), "赛事4次助攻，中场推进核心"),
        StarSpec("马丁内利", "桑巴猎豹", "silver", (94, 74, 88, 45, 90, 4), "高速纵向冲击与无球跑动"),
    ],
    "葡萄牙": [
        StarSpec("罗纳尔多", "葡国战神", "gold", (80, 90, 91, 42, 82, 5), "禁区终结、制空和关键球核心"),
        StarSpec("迪奥戈·科斯塔", "葡国门神", "silver", (62, 88, 86, 95, 90, 5), "赛事门将防守榜前列"),
        StarSpec("布鲁诺·费尔南德斯", "葡国司令", "silver", (82, 80, 92, 70, 91, 5), "直塞、远射和定位球主脑"),
        StarSpec("维蒂尼亚", "葡国节拍", "silver", (80, 75, 93, 76, 92, 4), "中场控球与压力下出球核心"),
    ],
    "德国": [
        StarSpec("诺伊尔", "战车门神", "gold", (56, 86, 88, 94, 86, 5), "清道夫门将；出击、指挥和大赛经验核心"),
        StarSpec("穆西亚拉", "日耳魔术", "silver", (92, 76, 92, 50, 86, 5), "狭小空间持球与推进核心"),
        StarSpec("维尔茨", "莱茵画师", "silver", (86, 73, 92, 58, 90, 5), "赛事3次助攻，前场创造者"),
        StarSpec("基米希", "德意铁轴", "silver", (78, 80, 91, 88, 92, 5), "多位置组织与防守指挥"),
    ],
    "日本": [
        StarSpec("久保建英", "蓝武左刃", "gold", (88, 68, 88, 45, 86, 5), "右侧持球、内切与创造力核心"),
        StarSpec("堂安律", "蓝武强弓", "silver", (86, 70, 85, 55, 88, 4), "边路内切和远射威胁"),
        StarSpec("铃木彩艳", "蓝武门神", "silver", (60, 85, 79, 88, 86, 4), "反应与覆盖范围突出的年轻门将"),
        StarSpec("镰田大地", "蓝武棋手", "silver", (78, 76, 85, 68, 87, 4), "前腰串联与第二落点处理"),
    ],
    "摩洛哥": [
        StarSpec("哈基米", "沙漠飞翼", "gold", (94, 80, 88, 88, 94, 5), "右路推进、回追和反击发动机"),
        StarSpec("布努", "北非门神", "silver", (58, 86, 82, 92, 88, 5), "淘汰赛关键扑救与大赛经验"),
        StarSpec("卜拉欣·迪亚斯", "北非魔术", "silver", (89, 72, 91, 48, 89, 5), "赛事4次助攻，前场创造核心"),
        StarSpec("赛巴里", "北非中枢", "silver", (82, 80, 87, 78, 91, 4), "中场推进和攻防转换关键点"),
    ],
    "挪威": [
        StarSpec("哈兰德", "北海魔人", "gold", (92, 98, 92, 42, 93, 5), "对巴西梅开二度；赛事7球"),
        StarSpec("厄德高", "北海司令", "silver", (84, 78, 93, 62, 92, 5), "赛事3次助攻，中场主脑"),
        StarSpec("尼兰", "北海铁壁", "silver", (56, 83, 76, 90, 87, 5), "淘汰赛对巴西多次关键扑救"),
    ],
    "哥伦比亚": [
        StarSpec("路易斯·迪亚斯", "咖啡飞翼", "gold", (93, 78, 88, 50, 91, 5), "边路推进与反击终结核心"),
        StarSpec("哈梅斯", "咖啡魔杖", "silver", (72, 74, 91, 58, 82, 5), "定位球、传中与关键传球主脑"),
        StarSpec("卡米罗·巴尔加斯", "咖啡门神", "silver", (55, 82, 73, 86, 84, 4), "稳定门线技术与大赛经验"),
    ],
    "美国": [
        StarSpec("普利希奇", "星条飞翼", "gold", (91, 72, 88, 45, 88, 5), "主场核心，边路持球与终结兼备"),
        StarSpec("麦肯尼", "星条铁腰", "silver", (80, 88, 84, 80, 92, 4), "中场对抗、前插和二点球核心"),
        StarSpec("安东尼·鲁宾逊", "星条快翼", "silver", (92, 82, 80, 84, 94, 4), "高强度左路往返与推进"),
    ],
    "加拿大": [
        StarSpec("阿方索·戴维斯", "枫叶闪电", "gold", (98, 82, 88, 82, 94, 5), "左路速度、推进和覆盖范围核心"),
        StarSpec("乔纳森·戴维", "枫叶猎手", "silver", (88, 80, 87, 45, 90, 4), "反击跑位与禁区终结核心"),
        StarSpec("科尼利厄斯", "枫叶铁塔", "silver", (76, 88, 70, 87, 88, 4), "后场对抗、制空与防线稳定器"),
        StarSpec("内森·萨利巴", "枫叶新星", "silver", (80, 78, 84, 70, 90, 4), "赛事2次助攻，中场推进新星"),
    ],
    "墨西哥": [
        StarSpec("基尼奥内斯", "绿鹰中锋", "gold", (90, 88, 86, 44, 91, 5), "纵向冲击、对抗和终结核心"),
        StarSpec("阿尔瓦雷斯", "绿鹰铁腰", "silver", (78, 90, 82, 88, 91, 4), "防守覆盖与攻防转换枢纽"),
        StarSpec("劳尔·希门尼斯", "绿鹰支点", "silver", (80, 88, 84, 45, 85, 4), "支点、制空和禁区经验"),
        StarSpec("阿尔瓦拉多", "绿鹰飞翼", "silver", (88, 74, 84, 48, 89, 4), "赛事3次助攻，边路创造者"),
    ],
    "佛得角": [
        StarSpec("沃齐尼亚", "蓝鲨门神", "gold", (54, 82, 70, 87, 84, 5), "门线反应、大赛经验和弱队上限核心"),
        StarSpec("瑞安·门德斯", "蓝鲨队长", "silver", (82, 78, 80, 45, 86, 5), "经验、定位球和反击终结核心"),
        StarSpec("凯文·皮纳", "蓝鲨铁腰", "silver", (72, 85, 75, 82, 88, 4), "中场对抗与防线保护"),
    ],
    "库拉索": [
        StarSpec("埃洛伊·鲁姆", "海岛门神", "gold", (54, 82, 72, 88, 84, 5), "赛事门将表现榜前列，弱队上限核心"),
        StarSpec("莱安多·巴库纳", "海岛司令", "silver", (76, 82, 82, 72, 87, 4), "队长、中场组织与定位球核心"),
        StarSpec("儒尼尼奥·巴库纳", "海岛节拍", "silver", (78, 78, 81, 68, 86, 4), "推进、持球和攻防串联"),
    ],
}

GENERIC_NAMES = {
    "GK": ["主力门将", "替补门将"],
    "DF": ["主力中卫", "替补中卫", "年轻中卫", "高大中卫", "速度中卫", "经验中卫", "左路边卫", "右路边卫", "进攻边卫", "防守边卫"],
    "MF": ["组织中场", "防守中场", "全能中场", "替补中场", "年轻中场", "经验中场", "技术中场", "跑动中场", "进攻中场", "主力前腰", "替补前腰", "防守后腰"],
    "FW": ["主力中锋", "替补中锋", "强力中锋", "速度边锋", "技术边锋", "替补边锋", "主力前锋", "年轻前锋", "替补前锋"],
}

BASE = {
    "GK": [56, 80, 68, 86, 80],
    "DF": [75, 80, 68, 83, 82],
    "MF": [76, 76, 80, 70, 84],
    "FW": [84, 76, 80, 42, 82],
}

MAINLAND_NAME_OVERRIDES = {
    "尼古拉斯·冈萨鲁伊斯": "尼古拉斯·冈萨雷斯",
    "哈梅斯·查福特": "詹姆斯·特拉福德",
    "连拿度·维加": "雷纳托·维加",
    "柏斯高·哥斯": "帕斯卡尔·格罗斯",
    "理查德·莱昂斯": "理查德·里奥斯",
    "阿利斯泰尔·約翰斯顿": "阿利斯泰尔·约翰斯顿",
    "尼可·西盖伊": "尼科·西古尔",
    "乔纳森·奥索莱昂": "乔纳森·奥索里奥",
    "埃泽雷尔·雷耶斯": "伊斯雷尔·雷耶斯",
    "布里安·古铁鲁伊斯": "布莱恩·古铁雷斯",
    "凯尔文·皮鲁伊斯": "凯尔文·皮雷斯",
    "泰鲁伊斯·诺斯林": "泰雷斯·诺斯林",
}


def clean_name(name: str) -> str:
    cleaned = re.sub(r"[（(][^）)]*(?:队长|副队长|captain)[^）)]*[）)]", "", name, flags=re.I).strip()
    return MAINLAND_NAME_OVERRIDES.get(cleaned, cleaned)


def parse_int(value: str) -> int:
    match = re.search(r"\d+", value)
    return int(match.group()) if match else 0


def player_age(cells: list[str]) -> int:
    match = re.search(r"（(\d+)岁）", cells[3])
    return int(match.group(1)) if match else 26


def star_for(team: str, name: str) -> StarSpec | None:
    matches = [spec for spec in STARS[team] if spec.keyword in name]
    if len(matches) > 1:
        raise ValueError(f"Ambiguous star mapping: {team} {name} {matches}")
    return matches[0] if matches else None


def jitter(team: str, name: str, index: int) -> int:
    digest = hashlib.sha256(f"{team}|{name}|{index}".encode("utf-8")).digest()
    return digest[index] % 7 - 3


def normal_stats(team: str, player: dict[str, str]) -> tuple[int, int, int, int, int, int]:
    position = POSITION_MAP[player["position"]]
    values = BASE[position].copy()
    caps = parse_int(player["cells"][4])
    goals = parse_int(player["cells"][5])
    age = player_age(player["cells"])
    experience = 2 if caps >= 70 else 1 if caps >= 35 else 0 if caps >= 12 else -1
    for index in range(5):
        values[index] += TEAM_META[team].delta + experience + jitter(team, player["name"], index)
    if age >= 33:
        values[0] -= 3
        values[2] += 1
        values[4] -= 2
    elif age <= 22:
        values[0] += 2
        values[4] += 1
    if position == "FW" and goals >= 20:
        values[2] += 2
    if position == "MF" and goals >= 10:
        values[2] += 1
    if position == "DF" and caps >= 60:
        values[3] += 2
    values = [max(28 if index == 3 and position == "FW" else 45, min(93, value)) for index, value in enumerate(values)]
    key = 4 if caps >= 65 or goals >= 20 else 3 if caps >= 10 else 2
    return (*values, key)


def power(position: str, stats: tuple[int, int, int, int, int, int]) -> float:
    spd, phy, tec, defense, sta, key = stats
    weights = {
        "GK": (0.05, 0.20, 0.15, 0.45, 0.15),
        "DF": (0.15, 0.20, 0.10, 0.40, 0.15),
        "MF": (0.12, 0.10, 0.30, 0.23, 0.25),
        "FW": (0.27, 0.17, 0.33, 0.05, 0.18),
    }[position]
    return sum(value * weight for value, weight in zip((spd, phy, tec, defense, sta), weights)) + key * 1.3


def price_for(position: str, stats: tuple[int, int, int, int, int, int], tier: str) -> int:
    premium = {"gold": 55, "silver": 25, "standard": 0}[tier]
    raw = (power(position, stats) - 54) * 3.25 + stats[5] * 11 + premium
    return max(40, min(260, int(round(raw / 5) * 5)))


def normal_description(position: str, stats: tuple[int, int, int, int, int, int], caps: int) -> str:
    labels = ["速度", "身体", "技术", "防守", "体能"]
    best = labels[max(range(5), key=lambda index: stats[index])]
    role = {"GK": "门线与出球轮换", "DF": "防线轮换", "MF": "中场轮换", "FW": "锋线轮换"}[position]
    experience = "经验丰富" if caps >= 60 else "具备国家队经验" if caps >= 20 else "阵容活力点"
    return f"{role}；{best}突出，{experience}"


def choose_alias(position: str, ordinal: int, used: set[str]) -> str:
    if not 1 <= ordinal <= len(GENERIC_NAMES[position]):
        raise RuntimeError(f"Generic name index out of range: {position} {ordinal}")
    candidate = GENERIC_NAMES[position][ordinal - 1]
    if candidate in used:
        raise RuntimeError(f"Duplicate generic alias: {candidate}")
    used.add(candidate)
    return candidate


def budget_for(players: list[dict]) -> int:
    descending = sorted((player["price"] for player in players), reverse=True)
    top13 = sum(descending[:13])
    top14 = sum(descending[:14])
    rounded = int(math.ceil(top13 / 10) * 10)
    return rounded if rounded < top14 else top13


def calibrate_prices(team: str, players: list[dict]) -> None:
    """Map each team's internal quality ranking onto the old plan's price ladder."""
    ladder = [240, 225, 210, 195, 180, 170, 160, 150, 140, 130, 120, 110, 100, 90, 80, 75, 70, 65, 60, 55, 50, 45, 40, 35]
    factor = {3: 1.08, 2: 1.03, 1: 1.00, 0: 0.95, -1: 0.92, -4: 0.78, -7: 0.68}[TEAM_META[team].delta]
    ranked = sorted(players, key=lambda player: player["price"], reverse=True)
    for player, base_price in zip(ranked, ladder):
        player["price"] = int(round(base_price * factor / 5) * 5)


def balanced_roster(players: list[dict], budget: int) -> tuple[list[dict], int]:
    """Find the strongest 18-player squad with real depth, not the 18 cheapest cards."""
    required_aliases = {player["alias"] for player in players if player["tier"] in ("gold", "silver")}
    best: tuple[float, int, list[dict]] | None = None
    for indexes in itertools.combinations(range(len(players)), 18):
        squad = [players[index] for index in indexes]
        if not required_aliases.issubset({player["alias"] for player in squad}):
            continue
        counts = {position: sum(player["position"] == position for player in squad) for position in ("GK", "DF", "MF", "FW")}
        if counts["GK"] != 2 or counts["DF"] < 5 or counts["MF"] < 5 or counts["FW"] < 3:
            continue
        cost = sum(player["price"] for player in squad)
        if cost > budget:
            continue
        quality = sum(power(player["position"], player["stats"]) for player in squad)
        candidate = (quality, cost, squad)
        if best is None or candidate[:2] > best[:2]:
            best = candidate
    if best is None:
        raise RuntimeError(f"No balanced 18-player roster under budget {budget}")
    return best[2], best[1]


def cheapest_count(players: list[dict], budget: int) -> int:
    spent = 0
    count = 0
    for price in sorted(player["price"] for player in players):
        if spent + price > budget:
            break
        spent += price
        count += 1
    return count


def build_data() -> dict[str, dict]:
    raw = parse()
    teams: dict[str, dict] = {}
    for team in TEAM_ORDER:
        players = []
        for source in raw[team]:
            name = clean_name(source["name"])
            if any(drop_name in name for drop_name in DROP[team]):
                continue
            source = dict(source)
            source["name"] = name
            players.append(source)
        if len(players) != 24:
            raise ValueError(f"{team}: expected 24, got {len(players)}")
        if sum(player["position"] == "门将" for player in players) != 2:
            raise ValueError(f"{team}: expected exactly 2 goalkeepers")

        reserved_star_aliases = {spec.alias for spec in STARS[team]}
        used = set(reserved_star_aliases)
        generic_counts = {"GK": 0, "DF": 0, "MF": 0, "FW": 0}
        resolved = []
        for source in players:
            position = POSITION_MAP[source["position"]]
            spec = star_for(team, source["name"])
            if spec:
                alias = spec.alias
                tier = spec.tier
                stats = spec.stats
                description = spec.note
            else:
                generic_counts[position] += 1
                alias = choose_alias(position, generic_counts[position], used)
                tier = "standard"
                stats = normal_stats(team, source)
                description = normal_description(position, stats, parse_int(source["cells"][4]))
            resolved.append({
                "number": int(source["number"]),
                "alias": alias,
                "reference": source["name"],
                "position": position,
                "tier": tier,
                "stats": stats,
                "price": price_for(position, stats, tier),
                "description": description,
            })
        if len({player["alias"] for player in resolved}) != 24:
            raise ValueError(f"{team}: public names must be unique within the team")
        if any(len(player["alias"]) != 4 or not all("\u3400" <= ch <= "\u9fff" for ch in player["alias"]) for player in resolved):
            raise ValueError(f"{team}: every public name must contain exactly four CJK characters")
        if sum(player["tier"] == "gold" for player in resolved) != 1:
            raise ValueError(f"{team}: gold card count is not 1")
        silver_count = sum(player["tier"] == "silver" for player in resolved)
        if silver_count not in (2, 3):
            raise ValueError(f"{team}: silver card count is {silver_count}")
        calibrate_prices(team, resolved)
        budget = budget_for(resolved)
        descending = sorted((player["price"] for player in resolved), reverse=True)
        if not (sum(descending[:13]) <= budget < sum(descending[:14])):
            raise ValueError(f"{team}: top-card recruitment calibration failed")
        balanced, balanced_cost = balanced_roster(resolved, budget)
        teams[team] = {
            "meta": TEAM_META[team],
            "players": sorted(resolved, key=lambda item: ("GKDFMFFW".index(item["position"]), item["number"])),
            "budget": budget,
            "total": sum(player["price"] for player in resolved),
            "balanced": balanced,
            "balanced_cost": balanced_cost,
            "cheapest_count": cheapest_count(resolved, budget),
        }
    return teams


INK = "17223B"
BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5B6573"
GOLD = "F6E6A8"
SILVER = "E3E7EC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=40, start=55, bottom=40, end=55) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_cell_width(cell, inches: float) -> None:
    width = int(inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_run_font(run, size: float, bold=False, color=INK, east_asia="Arial Unicode MS") -> None:
    run.font.name = "Arial Unicode MS"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Arial Unicode MS")
    r_fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    r_fonts.set(qn("w:eastAsia"), east_asia)


def add_text(paragraph, text: str, size=9.0, bold=False, color=INK):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, 8, color=MUTED)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(9)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 14, 7, BLUE),
        ("Heading 2", 12, 10, 5, BLUE),
        ("Heading 3", 10, 7, 4, INK),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.46)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    add_text(p, "剑指美加墨｜16队球员数据统一替换稿", 7.5, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    add_text(p, "第 ", 8, color=MUTED)
    add_field(p, "PAGE")
    add_text(p, " 页", 8, color=MUTED)


def add_title_block(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(5)
    add_text(p, "剑指美加墨", 13, bold=True, color=BLUE)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    add_text(p, "16队 × 24人球员数据与名单统一替换稿", 24, bold=True, color=INK)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(15)
    add_text(p, "旧策划书球员体系继承版｜2026-07-22", 10.5, color=MUTED)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(3)
        add_text(p, item, 9)


def add_rules(document: Document) -> None:
    document.add_heading("一、这次统一替换的最终口径", level=1)
    add_bullets(document, [
        "球队固定为16支：西班牙、阿根廷、法国、英格兰、巴西、葡萄牙、德国、日本、摩洛哥、挪威、哥伦比亚、美国、加拿大、墨西哥、佛得角、库拉索；新西兰不进入本批。",
        "每队固定24名候选球员，其中恰好2名门将；玩家从候选池中购买至少11人，不再采用“38人候选、征召23人”。每场仍选11人首发，并可从已购球员中设置最多5名替补。",
        "预算校准继承旧策划书并按本次口径修正：高价球星型阵容约13人；均衡型阵容固定按18人验算，必须包含2名门将、至少5后卫、5中场、3前锋，并带齐本队金银卡。纯低价路线可以超过18人，但阵容上限明显更低。",
        "每队1张金卡、2–3张银卡，其余普通卡。卡级不等于公开“总评”：静态表现继续由速度、身体、技术、防守、体能、关键时刻和价格决定；每场状态仍在60–100间动态生成并作为乘数。",
        "所有对外球员名均为四字中文名：普通卡直接使用“主力中卫、速度边锋、技术中场、替补门将”这类大众可理解的位置名；只有金卡、银卡保留个性化称号。普通位置名可跨球队复用，唯一性由teamId + playerId保证。",
    ])

    document.add_heading("二、字段与实现映射", level=1)
    rows = [
        ("公开字段", "建议数据键", "规则"),
        ("四字球员名", "name", "普通卡使用直白位置名；金银卡可个性化；同队内不重复"),
        ("现实关联", "referenceName", "仅内部维护；不在正式卡面默认展示"),
        ("号码 / 位置", "number / position", "位置使用GK、DF、MF、FW；每队2名GK"),
        ("卡级", "cardTier", "gold / silver / standard；旧版isGolden可临时由cardTier === 'gold'派生"),
        ("五项基础值", "spd / phy / tec / def / sta", "0–99；延续旧策划书算法与页面字段"),
        ("关键时刻", "star", "1–5星；参与关键节点成功率，不直接等同卡级"),
        ("价格", "price", "分；由位置权重、基础值、关键时刻和卡级溢价统一生成"),
        ("状态", "form", "不写死在名单；每场赛前60–100随机，连续首发继续累积疲劳"),
        ("描述", "description", "公开风格描述；金银卡可包含隐藏能力提示"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, text in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        add_text(p, text, 8.5, bold=True)
    for data in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(data):
            p = cells[idx].paragraphs[0]
            add_text(p, text, 8.2)
    widths = [1.15, 1.9, 7.0]
    set_table_fixed(table, widths)
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell, 55, 70, 55, 70)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_methodology(document: Document) -> None:
    document.add_heading("三、2026赛会表现调整原则", level=1)
    add_bullets(document, [
        "赛会奖项直接形成最高优先级修正：罗德里、梅西、姆巴佩、贝林厄姆、乌奈·西蒙、库巴西等按FIFA赛后奖项上调关键时刻及对应核心属性。",
        "助攻榜形成创造力修正：奥利塞5次助攻；吉马良斯、卜拉欣·迪亚斯4次；维尔茨、厄德高、阿尔瓦拉多等3次；这些球员重点上调技术与关键时刻。",
        "淘汰赛阶段形成球队层级修正：冠军/亚军/四强整体稳定性更高；八强黑马获得体能、防守或关键时刻补偿；较早出局球队保留个人球星上限，但普通轮换不整体虚高。",
        "非金银卡球员不逐人臆造赛会事件：其数值由球队层级、位置模板、国家队经验、年龄曲线和少量确定性差异生成；后续若有更细比赛数据，只需定点覆盖对应球员。",
    ])
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    add_text(p, "数据源：", 8.5, bold=True, color=BLUE)
    add_text(p, f"正式名单参考 {ROSTER_URL}；奖项 {AWARDS_URL}；助攻榜 {ASSISTS_URL}；淘汰赛赛果 {RESULTS_URL}；FIFA表现榜 {POWER_URL}。", 7.8, color=MUTED)


def tier_label(tier: str) -> str:
    return {"gold": "金", "silver": "银", "standard": "普"}[tier]


def stars_label(value: int) -> str:
    return "★" * value + "☆" * (5 - value)


def add_team_overview(document: Document, teams: dict[str, dict]) -> None:
    document.add_page_break()
    document.add_heading("四、16队总览", level=1)
    headers = ["球队", "赛会阶段", "金卡（现实关联）", "银卡", "24人总价", "预算", "均衡18成本", "低价最多"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, header, 7.2, bold=True)
    for team in TEAM_ORDER:
        data = teams[team]
        gold = next(player for player in data["players"] if player["tier"] == "gold")
        silvers = [player["alias"] for player in data["players"] if player["tier"] == "silver"]
        values = [
            team, data["meta"].result, f"{gold['alias']}（{gold['reference']}）", "、".join(silvers),
            str(data["total"]), str(data["budget"]), str(data["balanced_cost"]), str(data["cheapest_count"]),
        ]
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, 4, 5, 6, 7) else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, 6.8, bold=(idx == 0))
    widths = [0.58, 0.65, 1.85, 3.75, 0.70, 0.66, 0.78, 0.68]
    set_table_fixed(table, widths)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(19)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_team_section(document: Document, team: str, data: dict, index: int) -> None:
    meta = data["meta"]
    p = document.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(3)
    add_text(p, f"{index:02d}  {team}", 16, bold=True, color=BLUE)
    add_text(p, f"    {meta.difficulty}｜2026：{meta.result}｜预算：{data['budget']}分｜24人总价：{data['total']}分", 9, color=MUTED)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    add_text(p, f"球衣：{meta.jersey}　｜　球队技能：【{meta.skill}】　｜　队伍性格：{meta.personality}", 8.2)

    headers = ["#", "四字名", "现实关联（内部）", "位", "卡", "SPD", "PHY", "TEC", "DEF", "STA", "关键", "价格", "描述 / 赛会依据"]
    widths = [0.30, 0.67, 1.28, 0.34, 0.31, 0.34, 0.34, 0.34, 0.34, 0.34, 0.47, 0.43, 4.00]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, header, 6.7, bold=True)
    set_repeat_table_header(table.rows[0])
    for player in data["players"]:
        spd, phy, tec, defense, sta, key = player["stats"]
        values = [
            str(player["number"]), player["alias"], player["reference"], player["position"], tier_label(player["tier"]),
            str(spd), str(phy), str(tec), str(defense), str(sta), stars_label(key), str(player["price"]), player["description"],
        ]
        cells = table.add_row().cells
        fill = GOLD if player["tier"] == "gold" else SILVER if player["tier"] == "silver" else WHITE
        for idx, value in enumerate(values):
            if player["tier"] != "standard" and idx in (1, 4):
                set_cell_shading(cells[idx], fill)
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in (1, 2, 12) else WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, value, 6.35, bold=(idx == 1 or (idx == 4 and player["tier"] != "standard")))
        prevent_row_split(table.rows[-1])
    set_table_fixed(table, widths)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(15.2)
        for cell in row.cells:
            set_cell_margins(cell, 22, 35, 22, 35)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_document(teams: dict[str, dict]) -> Document:
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_title_block(document)
    add_rules(document)
    add_methodology(document)
    add_team_overview(document, teams)
    document.add_heading("五、16队完整24人数据", level=1)
    for index, team in enumerate(TEAM_ORDER, start=1):
        add_team_section(document, team, teams[team], index)
    final_heading = document.add_heading("六、落地替换检查清单", level=1)
    final_heading.paragraph_format.page_break_before = True
    add_bullets(document, [
        "先扩充cardTier并保留isGolden兼容派生；银卡只增加视觉层级和隐藏技能，不改旧页面的基础属性读取路径。",
        "导入384人时校验：每队24人、2名GK、1金、2–3银、四字名同队内不重复、号码在该队不重复、五项属性0–99、关键时刻1–5星。",
        "招募页改为至少购买11人即可继续；不设置固定23人上限。阵容页仍要求11名首发、首发1名GK；最多5名替补延续旧策划书。",
        "预算测试必须分别覆盖：高价球星型第14人被阻止；带齐金银卡、2门将和完整位置深度的均衡阵容可达到18人；纯低价路线允许超过18人但整体能力明显降低。",
        "公开界面默认只显示四字名；现实关联列进入内部数据或审计表，不直接出现在正式商业版本卡面。",
    ])
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_text(p, "文档结论：", 9.2, bold=True, color=BLUE)
    add_text(p, "本稿已经把旧策划书的球员字段、动态状态和预算体感完整迁移到16队×24人的新范围，可作为下一轮统一替换的唯一名单基准。", 9.2)
    return document


def main() -> None:
    teams = build_data()
    document = build_document(teams)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    for team in TEAM_ORDER:
        data = teams[team]
        gold = next(player["alias"] for player in data["players"] if player["tier"] == "gold")
        silvers = ",".join(player["alias"] for player in data["players"] if player["tier"] == "silver")
        print(f"{team}: players=24 GK=2 gold={gold} silver={silvers} total={data['total']} budget={data['budget']} balanced18={data['balanced_cost']} cheapMax={data['cheapest_count']}")
    print(OUTPUT)


if __name__ == "__main__":
    main()
