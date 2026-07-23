from pathlib import Path
import sys

from docx import Document


def main() -> None:
    path = Path(sys.argv[1])
    document = Document(path)
    print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)}")
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f"P{index:04d}\t{paragraph.style.name}\t{text}")
    for table_index, table in enumerate(document.tables):
        print(f"TABLE {table_index} rows={len(table.rows)} cols={len(table.columns)}")
        for row_index, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            print(f"R{row_index:04d}\t" + " | ".join(cells))


if __name__ == "__main__":
    main()
