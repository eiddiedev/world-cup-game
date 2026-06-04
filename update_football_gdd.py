from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path("/Users/a1234/Desktop/资产/剑指美加墨_完整游戏策划书.docx")
OUT = Path("/Users/a1234/Desktop/资产/剑指美加墨_完整游戏策划书_更新版.docx")

GOLD = "FFD966"
GOLD_DARK = "9C6500"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color


def format_gold_row(row):
    for cell in row.cells:
        shade_cell(cell, GOLD)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True if cell is row.cells[0] else run.bold


def paragraph_with_text(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def replace_paragraph(doc, old, new):
    paragraph = paragraph_with_text(doc, old)
    paragraph.text = new
    return paragraph


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph._element = new_p
    if style:
        new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def insert_table_after(paragraph, rows, cols):
    doc = paragraph._parent
    table = doc.add_table(rows=rows, cols=cols, width=Inches(6.5))
    tbl = table._tbl
    paragraph._p.addnext(tbl)
    return table


def append_table_after(paragraph, headers, rows, style_name="Table Grid"):
    table = insert_table_after(paragraph, len(rows) + 1, len(headers))
    try:
        table.style = style_name
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "F2F2F2")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            set_cell_text(table.rows[r_idx].cells[c_idx], value)
    return table


def update_row(table, row_idx, values):
    row = table.rows[row_idx]
    for col_idx, value in values.items():
        set_cell_text(row.cells[col_idx], value)
    format_gold_row(row)


def add_asset_rows(table):
    existing = {(row.cells[0].text, row.cells[1].text, row.cells[2].text, row.cells[3].text) for row in table.rows}
    additions = [
        ("球员像素资产（已制作）", "64×64px透明底", "每队16张左右", "2名门将(gk/gk2)+1名单独命名球星+通用slice球员；阿根廷多余切片可作彩蛋/轮换"),
        ("球星金卡卡面（已制作）", "64×64px角色像+金色卡底", "10张", "对应姆巴佩、内马尔、梅西、C罗、久保建英、哈兰德、阿什拉夫、诺伊尔、陈达毅、克里斯伍德"),
        ("棋盘对战表现（需制作UI）", "足球场俯视棋盘", "1套", "双方各11枚圆形号码棋子，只在关键节点展示跑位、传球、防守压迫和射门路线"),
    ]
    for addition in additions:
        if addition in existing:
            continue
        cells = table.add_row().cells
        for idx, value in enumerate(addition):
            set_cell_text(cells[idx], value)


def original_table(doc, original_idx):
    # The inserted "gold card" table is placed before the original formation
    # table, so all original table indexes from 2 onward shift by one.
    return doc.tables[original_idx + 1 if original_idx >= 2 else original_idx]


def main():
    doc = Document(SRC)

    replacements = {
        "10支球队 · 每队24人 · 精致像素风 · 世界杯热度蹭到底":
            "10支球队 · 每队24人大名单 · 每队16张左右像素资产复用 · 金卡球星 · 棋盘关键时刻",
        "美术风格：精致像素风（64×64球员全身站姿像，透明底）":
            "美术风格：精致像素风（64×64球员全身站姿像，透明底）；已准备球队、国旗、属性、足球、球场、庆祝等像素资产",
        "核心体验：选国家队 → 预算组建24人大名单 → 挑选11人首发+替补 → 赛程征途 → 剑指冠军":
            "核心体验：选国家队 → 用预算组建24人大名单 → 在16张左右像素角色资产中复用生成阵容 → 挑选11人首发+替补 → 关键时刻棋盘对战 → 剑指冠军",
        "选择国家队  →  组建24人大名单（预算限制）  →  排兵布阵（11人首发+阵型）  →  小组赛×3场  →  16强→8强→4强→决赛  →  结局":
            "选择国家队  →  组建24人大名单（预算限制）  →  排兵布阵（11人首发+阵型）  →  关键节点以足球场棋盘演出  →  小组赛×3场  →  16强→8强→4强→决赛  →  结局",
        "关键设计：预算限制使玩家无法买满所有24人，必须取舍；每场比赛前球员状态变化，使首发选择每场都有意义。":
            "关键设计：预算限制使玩家无法买满所有24人，必须取舍；每场比赛前球员状态变化，使首发选择每场都有意义；球星金卡拥有隐藏属性，能在关键节点改变棋盘局势。",
        "每队24人大名单：2名门将 + 22名场上球员（前锋/中场/后卫）":
            "每队玩法数据仍为24人大名单：2名门将 + 22名场上球员（前锋/中场/后卫）",
        "比赛：关键事件节点制，每场3–5个决策节点，节点间文字播报推进":
            "比赛：关键事件节点制，每场3–5个决策节点；平时用比分/文字播报推进，关键时刻切到足球场棋盘演出",
        "对战算法：纯本地计算，无需联网":
            "对战算法：纯本地计算，无需联网；关键节点同步输出棋盘事件（持球棋子、跑位棋子、防守棋子、射门/传球路线）",
        "包体估算：240张64×64 PNG约720KB + JS/CSS约100KB，总计约1MB，远低于8MB限制":
            "包体估算：约160张64×64球员PNG + 国旗/属性/球场/足球/庆祝资产 + JS/CSS，预计仍远低于8MB限制",
        "球员卡片：状态值用颜色区分（绿色=高状态/红色=低状态），信息密度适中":
            "球员卡片：普通球员用队伍色卡底，球星使用金色卡片；状态值用颜色区分（绿色=高状态/红色=低状态），信息密度适中",
    }
    for old, new in replacements.items():
        replace_paragraph(doc, old, new)

    roster_anchor = paragraph_with_text(doc, "门将必须有1名在首发，可选1名替补门将")
    insert_paragraph_after(roster_anchor, "资产复用规则：每队当前以16张左右球员像素资产为目标，其中gk/gk2固定用于两名门将，单独命名角色固定用于球星金卡，其余slice_XX可在普通球员、替补和轮换位之间复用；复用时通过号码、球员名、位置、状态色和卡片边框区分。")
    insert_paragraph_after(roster_anchor, "金卡规则：姆巴佩、内马尔、梅西、C罗、久保建英、哈兰德、阿什拉夫、诺伊尔、陈达毅、克里斯伍德在选人卡片中使用金色底框，并拥有不直接展示数值的隐藏属性。")

    hidden_rows = [
        ("法国", "法国超跑.png", "姆巴佩", "FW", "终点冲刺：反击、身后球、单刀节点速度判定+18%，射门成功率+8%。"),
        ("巴西", "桑巴舞者.png", "内马尔", "FW/MF", "桑巴单挑：1v1盘带节点技术判定+20%，同时更容易制造任意球。"),
        ("阿根廷", "当世球王.png", "梅西", "FW/MF", "最后一传：禁区前沿节点技术+20%，成功后队友下一脚射门+10%。"),
        ("葡萄牙", "边路游龙.png", "C罗", "FW", "终局头槌：75分钟后传中、定位球、争顶节点身体+15%，关键时刻临时+1星。"),
        ("日本", "蓝武锋魂.png", "久保建英", "FW/MF", "小空间转身：高压逼抢和狭小空间节点技术+16%，丢球风险-8%。"),
        ("挪威", "北欧魔人.png", "哈兰德", "FW", "禁区引力：直塞、抢点、禁区射门节点身体+18%，射门成功率+15%。"),
        ("摩洛哥", "北非之狐.png", "阿什拉夫", "DF/MF", "右路弹射：边路推进速度+15，回追防守+12，成功后下一次传中+10%。"),
        ("德国", "战车门卫.png", "诺伊尔", "GK", "清道夫门将：对手直塞/单刀节点可提前出击，扑救判定+15；失败时风险更高。"),
        ("库拉索", "蓝浪飞翼.png", "陈达毅", "FW", "加勒比闪击：替补登场后20分钟速度+15%，点球大战胆量+10%。"),
        ("新西兰", "全白重炮.png", "克里斯伍德", "FW", "全白支点：传中、高空球、二点球节点身体+16%，队友补射成功率+12%。"),
    ]
    table_anchor = paragraph_with_text(doc, "换人机会：每场比赛最多3次换人")
    title = insert_paragraph_after(table_anchor, "4.3 球星金卡与隐藏属性（选人卡片金色展示）")
    title.runs[0].bold = True
    append_table_after(title, ["球队", "资产文件", "金卡球星", "位置", "隐藏属性"], hidden_rows)

    flow_anchor = paragraph_with_text(doc, "赛后：比分 + 战报摘要 + 球员状态更新（连续上场球员下场状态-5）")
    chess_title = insert_paragraph_after(flow_anchor, "6.2.1 足球场棋盘呈现")
    chess_title.runs[0].bold = True
    insert_paragraph_after(chess_title, "对战画面不做全程实时跑动，只在关键时刻出现俯视足球场棋盘：双方各11枚圆形号码棋子，玩家队伍与对手用不同颜色区分。系统根据阵型把棋子摆在固定站位上，再用2–4步短动画表现持球、跑位、压迫、传球或射门。")
    insert_paragraph_after(chess_title, "关键事件示例：1）边路推进：边锋棋子沿边线前移，边后卫重叠；2）中路渗透：10号位棋子前插，前锋拉开中卫；3）防守封堵：后腰棋子回撤形成三角包夹；4）点球/单刀：只保留射手棋子与门将棋子，强化选择方向和结果反馈。")

    player_note = paragraph_with_text(doc, "说明：每队24人（2门将+22场上球员）。价格为购买所需预算分数。状态(FOR)为赛前随机生成，策划书中不列固定值，仅展示基础六项属性供参考。所有球员名称均为风格化虚构名称，不使用真实球员姓名，规避版权风险。")
    player_note.text = "说明：每队玩法数据仍为24人（2门将+22场上球员），但美术资产按每队约16张复用；价格为购买所需预算分数。状态(FOR)为赛前随机生成。黑客松演示版球星金卡使用真实姓名；如正式商业上线，可按授权情况切换为风格化称号。"

    asset_replacements = {
        "10.1 球员全身像（已生成）": "10.1 球员全身像与复用规则（已生成）",
        "数量：每队24人 × 10队 = 240张": "数量：每队当前约16张球员像素资产；玩法层仍生成24人大名单，普通球员通过slice_XX复用补足",
        "命名建议：france_01_fw.png / france_gk01.png 等": "命名现状：国家文件夹/gk.png、gk2.png、slice_XX.png、单独命名球星.png；正式开发直接按现有文件名引用",
        "10.2 待制作资产": "10.2 已准备资产与仍需制作UI",
    }
    for old, new in asset_replacements.items():
        replace_paragraph(doc, old, new)

    star_updates = [
        (4, 3, {0: "姆巴佩", 8: "金卡球星。速度爆点和反击终结能力顶级，身后球节点最危险"}),
        (5, 3, {0: "姆巴佩", 4: "金卡球星。速度爆点和反击终结能力顶级，身后球节点最危险"}),
        (6, 3, {0: "内马尔", 8: "金卡球星。盘带、假动作和小空间创造力顶级，能制造任意球"}),
        (7, 3, {0: "内马尔", 4: "金卡球星。盘带、假动作和小空间创造力顶级，能制造任意球"}),
        (8, 3, {0: "梅西", 8: "金卡球星。禁区前沿的最后一传和终结都是决定比赛的答案"}),
        (9, 3, {0: "梅西", 4: "金卡球星。禁区前沿的最后一传和终结都是决定比赛的答案"}),
        (10, 3, {0: "C罗", 8: "金卡球星。冲击禁区、头球和关键时刻终结能力仍是最大武器"}),
        (11, 3, {0: "C罗", 4: "金卡球星。冲击禁区、头球和关键时刻终结能力仍是最大武器"}),
        (12, 1, {0: "诺伊尔", 7: "★★★★★", 8: "金卡球星。清道夫门将，能提前出击化解单刀和身后球"}),
        (13, 1, {0: "诺伊尔", 3: "★★★★★", 4: "金卡球星。清道夫门将，能提前出击化解单刀和身后球"}),
        (14, 10, {0: "久保建英", 8: "金卡球星。小空间转身和左脚创造力，是日本前场的破局点"}),
        (15, 10, {0: "久保建英", 4: "金卡球星。小空间转身和左脚创造力，是日本前场的破局点"}),
        (16, 3, {0: "哈兰德", 8: "金卡球星。禁区内身体、抢点和射门压迫感极强"}),
        (17, 3, {0: "哈兰德", 4: "金卡球星。禁区内身体、抢点和射门压迫感极强"}),
        (18, 18, {0: "阿什拉夫", 1: "DF/MF", 2: "92", 3: "78", 4: "78", 5: "86", 6: "88", 7: "★★★★★", 8: "金卡球星。右路高速推进和回追防守兼备，是反击发动机"}),
        (19, 18, {0: "阿什拉夫", 1: "DF/MF", 2: "185", 3: "★★★★★", 4: "金卡球星。右路高速推进和回追防守兼备，是反击发动机"}),
        (20, 3, {0: "克里斯伍德", 8: "金卡球星。禁区支点和高空球优势明显，弱队破局核心"}),
        (21, 3, {0: "克里斯伍德", 4: "金卡球星。禁区支点和高空球优势明显，弱队破局核心"}),
        (22, 3, {0: "陈达毅", 8: "金卡球星。边路爆发和替补冲击力突出，适合反击偷袭"}),
        (23, 3, {0: "陈达毅", 4: "金卡球星。边路爆发和替补冲击力突出，适合反击偷袭"}),
    ]
    for table_idx, row_idx, values in star_updates:
        update_row(original_table(doc, table_idx), row_idx, values)

    add_asset_rows(original_table(doc, 25))

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {"4.3 球星金卡与隐藏属性（选人卡片金色展示）", "6.2.1 足球场棋盘呈现"}:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
