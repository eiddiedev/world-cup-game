#!/usr/bin/env python3
"""Build the merged game design document from the retained desktop template."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path('/Users/a1234/Desktop/策划书.docx')
SOURCE = ROOT / 'docs' / '剑指美加墨_游戏策划书.md'
OUTPUT = ROOT / 'docs' / '剑指美加墨_游戏策划书_2026-07-26.docx'
DESKTOP_OUTPUT = Path('/Users/a1234/Desktop/剑指美加墨_游戏策划书_2026-07-26.docx')

BLUE = '1A56A0'
DEEP_BLUE = '1B3764'
RED_BROWN = 'B34235'
BROWN = '5B4630'
LIGHT_BLUE = 'EAF2F8'
PALE_BLUE = 'F5F8FC'
WHITE = 'FFFFFF'
BLACK = '222222'
MUTED = '666666'
TABLE_WIDTH = 9360


def set_font(run, size=None, bold=None, color=None, italic=None, family='Microsoft YaHei'):
    run.font.name = family
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn('w:ascii'), family)
    fonts.set(qn('w:hAnsi'), family)
    fonts.set(qn('w:eastAsia'), family)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def shade(element, fill):
    shd = element.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        element.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for edge, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_borders(table, color=BLUE, size='6'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = borders.find(qn(f'w:{edge}'))
        if tag is None:
            tag = OxmlElement(f'w:{edge}')
            borders.append(tag)
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), size)
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)


def remove_body_content(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    normal = doc.styles['Normal']
    normal.font.name = 'Microsoft YaHei'
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.18

    specs = {
        'Heading 1': (16, DEEP_BLUE, 20, 10),
        'Heading 2': (13, RED_BROWN, 15, 7.5),
        'Heading 3': (11.5, BROWN, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abstract_id))
    multi = OxmlElement('w:multiLevelType')
    multi.set(qn('w:val'), 'singleLevel')
    abstract.append(multi)
    level = OxmlElement('w:lvl')
    level.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    level.append(start)
    num_fmt = OxmlElement('w:numFmt')
    num_fmt.set(qn('w:val'), 'decimal' if kind == 'number' else 'bullet')
    level.append(num_fmt)
    lvl_text = OxmlElement('w:lvlText')
    lvl_text.set(qn('w:val'), '%1.' if kind == 'number' else '•')
    level.append(lvl_text)
    suffix = OxmlElement('w:suff')
    suffix.set(qn('w:val'), 'tab')
    level.append(suffix)
    p_pr = OxmlElement('w:pPr')
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'num')
    tab.set(qn('w:pos'), '540')
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '540')
    ind.set(qn('w:hanging'), '280')
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn('w:numPr'))
    if num_pr is None:
        num_pr = OxmlElement('w:numPr')
        p_pr.append(num_pr)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    num_ref = OxmlElement('w:numId')
    num_ref.set(qn('w:val'), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_ref)


def add_inline(paragraph, text, size=None, color=None):
    for part in re.split(r'(\*\*.*?\*\*|`.*?`)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True, color=color)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size or 9.5, color=color or BLACK, family='Arial Unicode MS')
            shade(run._element.get_or_add_rPr(), PALE_BLUE)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, color=color)


def parse_table(lines):
    rows = [[cell.strip() for cell in line.strip().strip('|').split('|')] for line in lines]
    if len(rows) > 1 and all(re.fullmatch(r':?-{3,}:?', cell) for cell in rows[1]):
        rows.pop(1)
    return rows


def choose_widths(rows):
    cols = max(len(row) for row in rows)
    lengths = []
    for col in range(cols):
        values = [len(row[col]) if col < len(row) else 0 for row in rows]
        lengths.append(max(5, min(42, max(values, default=5))))
    total = sum(lengths)
    widths = [max(900, int(TABLE_WIDTH * value / total)) for value in lengths]
    excess = sum(widths) - TABLE_WIDTH
    if excess > 0:
        for index in sorted(range(cols), key=lambda i: widths[i], reverse=True):
            reducible = max(0, widths[index] - 900)
            reduction = min(excess, reducible)
            widths[index] -= reduction
            excess -= reduction
            if excess <= 0:
                break
    if sum(widths) < TABLE_WIDTH:
        widths[-1] += TABLE_WIDTH - sum(widths)
    return widths


def add_table(doc, rows):
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = choose_widths(rows)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(TABLE_WIDTH))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '110')
    tbl_ind.set(qn('w:type'), 'dxa')

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)

    set_table_borders(table)
    font_size = 8.2 if cols >= 6 else 8.8 if cols == 5 else 9.2 if cols == 4 else 9.5
    for row_index, values in enumerate(rows):
        tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        cant_split.set(qn('w:val'), 'true')
        tr_pr.append(cant_split)
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell._tc.get_or_add_tcPr(), BLUE if row_index == 0 else (LIGHT_BLUE if row_index % 2 == 0 else WHITE))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index < 2 and len(values[col_index] if col_index < len(values) else '') < 18 else WD_ALIGN_PARAGRAPH.LEFT
            text = values[col_index] if col_index < len(values) else ''
            add_inline(paragraph, text, size=font_size, color=WHITE if row_index == 0 else BLACK)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
        if row_index == 0:
            repeat = OxmlElement('w:tblHeader')
            repeat.set(qn('w:val'), 'true')
            tr_pr.append(repeat)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_opening(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(12)
    set_font(p.add_run('⚽ 剑指美加墨'), size=36, bold=True, color=BLUE, family='Arial Unicode MS')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_font(p.add_run('Targeting USA · Canada · Mexico'), size=14, bold=True, color=BLUE, family='Arial')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run('游戏策划书'), size=15, bold=True, color=DEEP_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_font(p.add_run('版本日期：2026-07-26  ·  H5 网页 / 抖音互动空间'), size=9.5, color=MUTED)


def build():
    doc = Document(REFERENCE)
    remove_body_content(doc)
    configure_styles(doc)
    add_opening(doc)

    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith('## '))
    number_id = None
    bullet_id = None
    previous_was_numbered = False
    previous_was_bullet = False
    in_code = False
    code_lines = []
    index = start
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                shade(p._p.get_or_add_pPr(), PALE_BLUE)
                set_font(p.add_run('\n'.join(code_lines)), size=9, color=BLACK, family='Arial Unicode MS')
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            previous_was_numbered = False
            previous_was_bullet = False
            index += 1
            continue
        if stripped.startswith('|') and index + 1 < len(lines) and lines[index + 1].strip().startswith('|'):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith('|'):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, parse_table(table_lines))
            previous_was_numbered = False
            previous_was_bullet = False
            continue
        heading = re.match(r'^(#{2,4})\s+(.*)$', stripped)
        if heading:
            level = len(heading.group(1)) - 1
            p = doc.add_paragraph(style=f'Heading {level}')
            add_inline(p, heading.group(2))
            previous_was_numbered = False
            previous_was_bullet = False
            index += 1
            continue
        numbered = re.match(r'^\d+\.\s+(.*)$', stripped)
        if numbered:
            if not previous_was_numbered:
                number_id = add_numbering(doc, 'number')
            p = doc.add_paragraph()
            apply_numbering(p, number_id)
            add_inline(p, numbered.group(1))
            previous_was_numbered = True
            previous_was_bullet = False
            index += 1
            continue
        bullet = re.match(r'^-\s+(.*)$', stripped)
        if bullet:
            if not previous_was_bullet:
                bullet_id = add_numbering(doc, 'bullet')
            p = doc.add_paragraph()
            apply_numbering(p, bullet_id)
            add_inline(p, bullet.group(1))
            previous_was_numbered = False
            previous_was_bullet = True
            index += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, stripped)
        previous_was_numbered = False
        previous_was_bullet = False
        index += 1

    doc.core_properties.title = '《剑指美加墨》游戏策划书'
    doc.core_properties.subject = '国家队经营、实时比赛与机制算法'
    doc.core_properties.author = '剑指美加墨项目组'
    doc.core_properties.comments = '合并策划书与机制算法，更新至 2026-07-26 当前版本。'
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    shutil.copy2(OUTPUT, DESKTOP_OUTPUT)
    print(OUTPUT)
    print(DESKTOP_OUTPUT)


if __name__ == '__main__':
    build()
