from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


ROOT = Path("/Users/a1234/Documents/ZZhackerthon")
SOURCE = Path("/Users/a1234/Desktop/资产/剑指美加墨_完整游戏策划书_最终命名版.docx")
TARGETS = [
    SOURCE,
    ROOT / "assets/剑指美加墨_完整游戏策划书_最终命名版.docx",
    ROOT / "public/assets/剑指美加墨_完整游戏策划书_最终命名版.docx",
]


def set_doc_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(27, 55, 100)),
        ("Heading 2", 13, RGBColor(179, 66, 53)),
        ("Heading 3", 11.5, RGBColor(91, 70, 48)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.paragraphs[
        list(paragraph._parent._element).index(new_p)
    ]
    inserted.clear()
    if text:
        inserted.add_run(text)
    if style:
        inserted.style = style
    return inserted


def add_after(anchor, blocks):
    current = anchor
    for text, style in reversed(blocks):
        current = insert_after(current, text, style)
    return current


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def replace_contains(doc: Document, needle: str, new_text: str) -> None:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            paragraph.text = new_text
            return
    raise ValueError(f"Paragraph containing not found: {needle}")


def color_lead(paragraph, color=RGBColor(27, 55, 100)) -> None:
    if paragraph.runs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.color.rgb = color


def build_front_matter() -> list[tuple[str, str | None]]:
    return [
        ("版本同步：2026-06-01 黑客松可玩演示版", "Heading 1"),
        (
            "一句话定位：一款把“世界杯贴纸册的收集冲动”与“GBA/FC式足球关键时刻棋盘演出”结合起来的轻量策略足球游戏。玩家不是单纯看比分，而是在预算、状态、阵型、换人、犯规和关键决策中，亲手把一支国家队推向美加墨世界杯冠军。",
            None,
        ),
        ("当前可玩闭环", "Heading 2"),
        (
            "已完成从首页、国家队选择、预算招募、排兵布阵、赛程推进、比赛模拟、关键决策、点球大战到赛后结算的完整闭环。演示版支持10支国家队、每队24人大名单、金卡球星、10种阵型、28类关键决策、红黄牌/伤停/换人事件、赛后状态变化与淘汰赛晋级。",
            None,
        ),
        ("创意亮点", "Heading 2"),
        (
            "1. 世界杯贴纸册式收集：国旗、金卡球星、像素球员和预算选择形成“想集齐但买不满”的心理张力。",
            None,
        ),
        (
            "2. FC式棋盘比赛：比赛不是纯文本模拟，22枚圆形号码棋子会在俯视球场上持续跑位；播报中的本方/对方号码、传球、封堵、犯规、扑救和射门会同步驱动画面。",
            None,
        ),
        (
            "3. 关键时刻可玩：玩家在单刀、任意球、点球、禁区混战、防守危机、换人危机等场景里做选择，每次选择都对应风险、收益、球员属性和动画结果。",
            None,
        ),
        (
            "4. 弱队也有故事：法国、巴西等强队更容易夺冠，但库拉索、新西兰等弱队通过低预算阵容、状态波动和关键决策也能制造戏剧性爆冷。",
            None,
        ),
        ("机制合理性", "Heading 2"),
        (
            "预算系统经过校准：每队24人大名单无法买满；全选强者约13人，中等取舍约15人，全选低价轮换约18人。这个区间保证玩家至少能凑出首发，又必须在球星、替补深度、门将和位置覆盖之间取舍。",
            None,
        ),
        (
            "比赛结果由球队强度、球员属性、阵型、状态、对手强度、关键决策和随机因子共同决定。强队拥有更稳定的上限，弱队有更高的风险和更强的叙事张力；红黄牌、伤停和体能衰减会迫使玩家在长期赛程中做轮换。",
            None,
        ),
        ("测试与验证", "Heading 2"),
        (
            "自动化回归：当前项目包含22项回归测试，覆盖主页路由、设置开关、旧存档兼容、音效播放路径、球队/球员数据、28个关键决策、动画模板完整性、任意球/点球专属动画、播报可视化事件、足球素材引用和决策触发间隔。",
            None,
        ),
        (
            "运行验证：最近一次开发验证已通过 npm test -- --run、npm run lint、npm run build，确保功能逻辑、代码规范和生产构建均可用。",
            None,
        ),
        (
            "平衡性验证：基于300次模拟，法国冠军率约41%、巴西38.7%、阿根廷25.3%、日本7.7%、库拉索0%。该结果体现强队优势明显、弱队不会异常夺冠，同时淘汰赛保留足球比赛应有的不确定性。",
            None,
        ),
        ("当前实现的沉浸反馈", "Heading 2"),
        (
            "音效与设置已接入：按钮反馈、开场哨、足球入网、庆祝、门将扑救、失球、红黄牌、换人、胜负结算均使用Web Audio合成；设置页可独立开关音效、音乐和震动。背景音乐与音效使用独立音频上下文，避免互相覆盖。",
            None,
        ),
        (
            "视觉风格已统一为世界杯贴纸册 × GBA游戏 × 开罗游戏 × Retro Bowl方向：像素字体、硬边框、硬阴影、米黄色/深蓝/红色/金色/棕色配色、移动端一屏信息优先。",
            None,
        ),
    ]


def append_development_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("附录：黑客松演示版开发同步", level=1)

    sections = [
        (
            "A. 已完成模块",
            [
                "首页：像素主视觉、开始/继续/设置入口、移动端适配。",
                "国家队选择：10支国家队、难度、预算、国旗、球队技能。",
                "球员招募：24人大名单、预算购买/出售、位置筛选、金卡球星、像素头像复用。",
                "排兵布阵：10种阵型、11人首发、替补席、位置适配与状态检查。",
                "赛程与晋级：小组赛、淘汰赛、具体国家对手解析、赛后推进。",
                "比赛：90分钟推进、连续播报、28类关键决策、FC式棋盘动画、红黄牌/伤停/换人。",
                "赛后：比分、数据、状态变化、伤停/红牌汇总、下一场或结局跳转。",
            ],
        ),
        (
            "B. 比赛表现层",
            [
                "22个圆形号码棋子按阵型锚点持续跑位，不再是静态摆拍。",
                "播报事件会生成visual payload，同一事件同时驱动文字和动画。",
                "决策前会从当前持球区域过渡到关键球员脚下，减少瞬移感。",
                "任意球具备人墙、主罚、传中/打门/重组织；点球具备主罚、门将扑救和方向结果。",
                "进球、扑救、红黄牌等结果具有同步音效和画面提示。",
            ],
        ),
        (
            "C. 评委可观察亮点",
            [
                "几分钟即可理解：选队、买人、排阵、开赛、关键时刻决策。",
                "每局都有变量：预算、球员状态、伤停红牌、换人、点球、淘汰赛压力。",
                "视觉记忆点明确：像素世界杯、金卡球星、贴纸册式国家队收集。",
                "适合互动空间：规则轻、反馈快、每次选择都能在画面上“被玩到”。",
            ],
        ),
    ]

    for heading, items in sections:
        doc.add_heading(heading, level=2)
        for item in items:
            p = doc.add_paragraph(style="List Paragraph")
            p.add_run("• " + item)


def update_doc(path: Path) -> None:
    doc = Document(path)
    set_doc_style(doc)

    subtitle = find_paragraph(doc, "完整游戏开发策划书 · 正式版")
    title_sync = insert_after(subtitle, "版本状态：黑客松可玩演示版 · 机制/视觉/测试同步更新", None)
    color_lead(title_sync, RGBColor(179, 66, 53))

    overview_anchor = find_paragraph(doc, "核心体验：选国家队 → 用预算组建24人大名单 → 在16张左右像素角色资产中复用生成阵容 → 挑选11人首发+替补 → 关键时刻棋盘对战 → 剑指冠军")
    add_after(overview_anchor, build_front_matter())

    replace_contains(
        doc,
        "比赛：关键事件节点制",
        "比赛：90分钟连续推进，平时以号码播报与足球场棋盘持续跑位呈现；每场约4–9个关键决策节点，强队因进攻推进和控场能力更容易制造关键时刻，弱队则更依赖少数高收益选择。",
    )
    replace_contains(
        doc,
        "关键事件示例：1）边路推进",
        "关键事件示例：1）边路推进：本方11号沿边线前插，边后卫重叠套上；2）中路渗透：10号位回撤接球，前锋拉开中卫；3）防守封堵：对方7号推进，本方4号补位封堵；4）任意球：人墙站位、主罚助跑、直接打门或传中争顶；5）点球：主罚球员、门将扑救方向和球路结果同步演出。",
    )
    replace_contains(
        doc,
        "对战画面不做全程实时跑动",
        "对战画面采用“阵型锚点 + 持球区域 + 事件脚本 + 结果动画”的FC式表现层：双方各11枚圆形号码棋子在俯视球场上持续轻微跑位，播报事件驱动传球、压迫、封堵、扑救和射门；关键决策前会从当前持球区域桥接到关键球员脚下，避免球权瞬移。",
    )
    replace_contains(
        doc,
        "黑客松演示版球星金卡使用真实姓名",
        "演示版金卡球星使用风格化称号，不直接使用真实姓名；金卡在招募卡片中使用金色底框并拥有隐藏风格加成。正式商业上线可继续沿用称号体系，避免授权风险。",
    )

    append_development_appendix(doc)
    doc.save(path)


def main() -> None:
    backup = SOURCE.with_name(SOURCE.stem + "_更新前备份.docx")
    if not backup.exists():
        shutil.copy2(SOURCE, backup)

    update_doc(SOURCE)
    for target in TARGETS[1:]:
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(SOURCE, target)


if __name__ == "__main__":
    main()
