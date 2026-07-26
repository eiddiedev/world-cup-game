#!/usr/bin/env python3
"""Build the two judge-facing DOCX files from their reviewed Markdown sources."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
OUTPUTS = [
    {
        "source": ROOT / "docs" / "JUDGE_GAME_DESIGN.md",
        "repo": ROOT / "docs" / "剑指美加墨_策划书_评委版_2026-07-26.docx",
        "desktop": DESKTOP / "剑指美加墨_策划书_评委版_2026-07-26.docx",
        "preset": "narrative_proposal",
        "cover": "proposal_centerpiece",
        "kicker": "抖音互动空间黑客松 · 评委材料",
        "subtitle": "从国家队经营到 11v11 实时比赛的完整足球叙事",
    },
    {
        "source": ROOT / "docs" / "MECHANICS_AND_BALANCE.md",
        "repo": ROOT / "docs" / "剑指美加墨_机制算法_评委版_2026-07-26.docx",
        "desktop": DESKTOP / "剑指美加墨_机制算法_评委版_2026-07-26.docx",
        "preset": "compact_reference_guide",
        "cover": "editorial_cover",
        "kicker": "可审计机制参考",
        "subtitle": "所有影响胜率的入口、公式、状态与验证证据",
    },
]

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
GOLD = "B8860B"
WHITE = "FFFFFF"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "B8C4D1"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_run_font(run, size=None, color=None, bold=None, italic=None, mono=False):
    # Named CJK compatibility override: the bundled LibreOffice renderer treats
    # Chinese glyphs as hAnsi in some runs, so use one Unicode font for every slot.
    western = "Arial Unicode MS"
    east_asia = "Arial Unicode MS"
    run.font.name = western
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(element, fill):
    shd = element.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        element.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    set_run_font(run, size=9, color=MUTED)


def create_numbering_instance(doc, start=1):
    numbering = doc.part.numbering_part.element
    existing_num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId"))
    ]
    num_id = max(existing_num_ids or [0]) + 1
    base_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    base_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
    num_pr.get_or_add_numId().set(qn("w:val"), str(num_id))


def configure_document(doc, preset, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8 if preset == "narrative_proposal" else 6)
    normal.paragraph_format.line_spacing = 1.333 if preset == "narrative_proposal" else 1.25
    normal.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if preset == "narrative_proposal" else WD_ALIGN_PARAGRAPH.LEFT
    )

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12 if preset == "narrative_proposal" else 14, 6 if preset == "narrative_proposal" else 7),
        "Heading 3": (12, DARK_BLUE, 8 if preset == "narrative_proposal" else 10, 4 if preset == "narrative_proposal" else 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194 if preset == "narrative_proposal" else -0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208 if preset == "narrative_proposal" else 1.25

    if "Code Block" not in [style.name for style in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["Code Block"]
    code_style.font.name = "Arial Unicode MS"
    code_style.font.size = Pt(9)
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.15
    code_style.paragraph_format.keep_together = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(running_title)
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = fp.add_run("剑指美加墨  ·  2026-07-26  ·  ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_field(fp)


def add_inline_runs(paragraph, text, size=None, color=None):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size or 9.5, color=INK, mono=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_cover(doc, title, subtitle, kicker, cover):
    if cover == "proposal_centerpiece":
        for _ in range(3):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        run = p.add_run(kicker)
        set_run_font(run, size=11, color=GOLD, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(title)
        set_run_font(run, size=26, color=INK, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(30)
        run = p.add_run(subtitle)
        set_run_font(run, size=14, color=DARK_BLUE)
        table = doc.add_table(rows=2, cols=2)
        set_table_geometry(table, [4680, 4680])
        values = [
            ("正式内容", "16 支可执教国家队 · 每队 24 人"),
            ("核心体验", "经营 · 情报 · 布阵 · 实时比赛 · 复盘 · 收集"),
            ("评审版本", "2026-07-26"),
            ("发布方向", "完整正式版 + 互动空间轻量包"),
        ]
        for cell, (label, value) in zip([c for row in table.rows for c in row.cells], values):
            shade(cell._tc.get_or_add_tcPr(), LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{label}\n")
            set_run_font(r, size=8.5, color=MUTED, bold=True)
            r = p.add_run(value)
            set_run_font(r, size=10.5, color=INK, bold=True)
    else:
        for _ in range(5):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(kicker)
        set_run_font(run, size=10.5, color=GOLD, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(title)
        set_run_font(run, size=28, color=INK, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(70)
        run = p.add_run(subtitle)
        set_run_font(run, size=14, color=DARK_BLUE)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("公式可追溯 · 状态可验证 · 结论可复现")
        set_run_font(run, size=11, color=MUTED, italic=True)
    doc.add_page_break()


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows


def widths_for(column_count):
    patterns = {
        2: [2400, 6960],
        3: [1900, 3600, 3860],
        4: [1550, 2450, 3150, 2210],
        5: [1450, 1850, 2700, 1600, 1760],
        6: [1400, 1400, 1600, 1600, 1600, 1760],
    }
    return patterns.get(column_count, [TABLE_WIDTH // column_count] * (column_count - 1) + [TABLE_WIDTH - (TABLE_WIDTH // column_count) * (column_count - 1)])


def add_table(doc, rows, preset):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths_for(column_count))
    header_fill = LIGHT_GRAY if preset == "narrative_proposal" else LIGHT_BLUE
    for row_index, values in enumerate(normalized):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                shade(cell._tc.get_or_add_tcPr(), header_fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_inline_runs(p, value, size=8.5 if column_count >= 4 else 9)
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(INK)
        if row_index == 0:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def markdown_to_docx(source, output, preset, cover, kicker, subtitle):
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = lines[0].lstrip("# ").strip()
    doc = Document()
    configure_document(doc, preset, title)
    add_cover(doc, title, subtitle, kicker, cover)

    index = 1
    in_code = False
    code_lines = []
    active_numbering_id = None
    previous_was_number = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                shade(p._p.get_or_add_pPr(), LIGHT_GRAY)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, size=9, color=INK, mono=True)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            previous_was_number = False
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, parse_table(table_lines), preset)
            previous_was_number = False
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, heading.group(2))
            index += 1
            previous_was_number = False
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, bullet.group(1))
            index += 1
            previous_was_number = False
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            requested_start = int(numbered.group(1))
            if not previous_was_number or requested_start == 1:
                active_numbering_id = create_numbering_instance(doc, requested_start)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, active_numbering_id)
            add_inline_runs(p, numbered.group(2))
            index += 1
            previous_was_number = True
            continue
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        index += 1
        previous_was_number = False

    doc.core_properties.title = title
    doc.core_properties.subject = "剑指美加墨 2026-07-26 评委材料"
    doc.core_properties.author = "剑指美加墨项目组"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    for item in OUTPUTS:
        markdown_to_docx(
            item["source"],
            item["repo"],
            item["preset"],
            item["cover"],
            item["kicker"],
            item["subtitle"],
        )
        shutil.copy2(item["repo"], item["desktop"])
        print(item["repo"])
        print(item["desktop"])


if __name__ == "__main__":
    main()
