from pathlib import Path

from docx import Document


SRC = Path("/Users/a1234/Desktop/资产/剑指美加墨_完整游戏策划书_预算校准版.docx")
OUT = Path("/Users/a1234/Desktop/资产/剑指美加墨_完整游戏策划书_最终命名版.docx")

NAME_MAP = {
    "姆巴佩": "法国超跑",
    "内马尔": "桑巴舞者",
    "梅西": "当世球王",
    "C罗": "边路游龙",
    "久保建英": "蓝武锋魂",
    "哈兰德": "北欧魔人",
    "阿什拉夫": "北非之狐",
    "诺伊尔": "战车门卫",
    "陈达毅": "蓝浪飞翼",
    "克里斯伍德": "全白重炮",
}


def replace_in_paragraph(paragraph):
    for run in paragraph.runs:
        text = run.text
        for old, new in NAME_MAP.items():
            text = text.replace(old, new)
        run.text = text


def replace_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph)


def main():
    doc = Document(SRC)
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        replace_in_table(table)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
