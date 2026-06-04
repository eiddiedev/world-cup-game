from pathlib import Path

from docx import Document


SRC = Path("/Users/a1234/Desktop/资产/剑指美加墨_完整游戏策划书_更新版.docx")
OUT = Path("/Users/a1234/Desktop/资产/剑指美加墨_完整游戏策划书_预算校准版.docx")

OLD_BUDGETS = {
    "法国": "1500",
    "巴西": "1500",
    "阿根廷": "1300",
    "葡萄牙": "1300",
    "德国": "1100",
    "日本": "1100",
    "挪威": "900",
    "摩洛哥": "900",
    "新西兰": "700",
    "库拉索": "700",
}

NEW_BUDGETS = {
    "法国": "2300",
    "巴西": "2250",
    "阿根廷": "2100",
    "葡萄牙": "2050",
    "德国": "1950",
    "日本": "1850",
    "挪威": "1700",
    "摩洛哥": "1800",
    "新西兰": "1280",
    "库拉索": "1170",
}

TOTALS = {
    "法国": "3305",
    "巴西": "3262",
    "阿根廷": "3043",
    "葡萄牙": "3016",
    "德国": "2854",
    "日本": "2678",
    "挪威": "2449",
    "摩洛哥": "2602",
    "新西兰": "1856",
    "库拉索": "1702",
}


def set_cell_text(cell, text):
    cell.text = ""
    cell.paragraphs[0].add_run(str(text))


def insert_after(paragraph, text):
    new_paragraph = paragraph.insert_paragraph_before(text)
    paragraph._p.addprevious(new_paragraph._p)
    return new_paragraph


def main():
    doc = Document(SRC)

    # Team overview table.
    overview = doc.tables[0]
    for row in overview.rows[1:]:
        team = row.cells[1].text.strip()
        if team in NEW_BUDGETS:
            set_cell_text(row.cells[2], NEW_BUDGETS[team])

    for paragraph in doc.paragraphs:
        text = paragraph.text
        for team, new_budget in NEW_BUDGETS.items():
            old_budget = OLD_BUDGETS[team]
            if text.startswith(f"{team}  ") and f"预算：{old_budget}分" in text:
                text = text.replace(f"预算：{old_budget}分", f"预算：{new_budget}分")
            total = TOTALS[team]
            old_line = f"24人总价：{total}分（远超预算{old_budget}分，玩家必须取舍，这正是乐趣所在）"
            new_line = (
                f"24人总价：{total}分；预算：{new_budget}分。经济体感：全选高价核心约13人，"
                "普通取舍约15人，全选低价轮换约18人，仍无法买满24人。"
            )
            if old_line in text:
                text = text.replace(old_line, new_line)
            if team == "摩洛哥" and "24人总价：2529分（远超预算900分，玩家必须取舍，这正是乐趣所在）" in text:
                text = text.replace(
                    "24人总价：2529分（远超预算900分，玩家必须取舍，这正是乐趣所在）",
                    "24人总价：2602分；预算：1800分。经济体感：全选高价核心约13人，普通取舍约15人，全选低价轮换约18人，仍无法买满24人。",
                )
        if text != paragraph.text:
            paragraph.text = text

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "关键设计：预算限制使玩家无法买满所有24人，必须取舍；每场比赛前球员状态变化，使首发选择每场都有意义；球星金卡拥有隐藏属性，能在关键节点改变棋盘局势。":
            note = paragraph.insert_paragraph_before(
                "预算校准目标：玩家从24人大名单中购买球员时，如果每个位置都选最强，大约只能买13人；如果在主力和替补之间做取舍，大约买15人；如果几乎全选低价球员，大约买18人。这样既能保证至少11人首发，也让替补、状态轮换和金卡选择形成真实策略。"
            )
            paragraph._p.addnext(note._p)
            break

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
